import json
from pathlib import Path
from typing import Optional
from datetime import datetime, timezone
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.config import settings

REPORT_PROMPT = """You are an M&A due diligence report writer specializing in INDIAN M&A transactions. Generate a comprehensive due diligence report based on the analysis below.

Deal: {deal_name}
Target: {target_company} (Indian company)
Acquirer: {acquirer}

Write the report considering the following INDIAN legal and regulatory framework:
- Companies Act 2013 (corporate governance, statutory compliance, board resolutions, RPTs)
- SEBI regulations (SAST Takeover Code, LODR, ICDR, PIT)
- RBI/FEMA (foreign investment, FDI policy, pricing guidelines, repatriation)
- CCI (Competition Act — merger control, combinations)
- Income Tax Act 1961 (TDS, transfer pricing, MAT, capital gains)
- CGST Act 2017 (GST compliance, input tax credit)
- DPDP Act 2023 (data protection, data localization)
- Indian labour laws (PF, ESI, Gratuity, Bonus, POSH, Contract Labour)
- Indian Contract Act 1872
- Arbitration and Conciliation Act 1996
- Limitation Act 1963

Document Summary:
{analysis_summary}

Issues Found:
{issue_summary}

Structure the report as follows:

# Due Diligence Report

## 1. Executive Summary
Brief overview of the target company and key findings

## 2. Corporate Structure & Governance (Companies Act 2013)
- Corporate structure, shareholding, board composition
- Statutory compliance (board meetings, annual filings, RPTs)
- CSR compliance (Section 135)
- Material contracts and shareholder agreements

## 3. Financial & Tax Due Diligence
- Financial health, revenue trends, profitability
- Tax compliance (TDS, GST, transfer pricing, MAT)
- Tax exposures, assessments, reassessment risk
- Tax holidays and benefits

## 4. Regulatory Compliance
- SEBI/RBI/FEMA compliance
- CCI merger control implications
- Sector-specific regulatory approvals
- FDI policy compliance

## 5. Legal & Contractual
- Key contracts (SHA, JVA, NDA, supply agreements)
- Litigation and dispute status
- IP ownership and registrations
- Real estate and property matters

## 6. Employment & Labour
- Workforce structure and cost
- Statutory compliance (PF, ESI, Gratuity, Bonus, POSH, CLRA)
- ESOP and employee benefits
- Labour law exposure assessment

## 7. Data & Technology
- DPDP Act 2023 compliance
- Data localization and data transfer
- IT systems and cybersecurity
- Software licensing and IP

## 8. Key Risks & Issues
Summary of all high/medium severity issues

## 9. Recommendations
Actionable recommendations for the acquirer

## 10. Deal Considerations
Key considerations for transaction structuring (share vs asset purchase, stamp duty, tax implications, regulatory approvals, timeline)
"""


def generate_report_content(deal_name, target_company, acquirer,
                             analysis_summary, issue_summary, client=None) -> str:
    from openai import OpenAI
    if client is None:
        client = OpenAI(api_key=settings.openai_api_key, base_url=settings.openai_base_url)
    try:
        resp = client.chat.completions.create(
            model=settings.openai_model,
            messages=[
                {"role": "system", "content": "You are an M&A due diligence report writer specialized in Indian M&A transactions."},
                {"role": "user", "content": REPORT_PROMPT.format(
                    deal_name=deal_name,
                    target_company=target_company,
                    acquirer=acquirer,
                    analysis_summary=analysis_summary[:12000],
                    issue_summary=issue_summary[:8000],
                )},
            ],
            temperature=0.3,
            max_tokens=5000,
        )
        return resp.choices[0].message.content
    except Exception as e:
        return f"Report generation failed: {str(e)}"


def generate_docx_report(report_text: str, output_path: str, deal_name: str) -> str:
    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    title = doc.add_heading(f"Due Diligence Report: {deal_name}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%d %B %Y')}")
    doc.add_paragraph("")

    for line in report_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line[0].isdigit() and ". " in line[:4]:
            doc.add_paragraph(line, style="List Number")
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(4)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path
